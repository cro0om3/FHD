import requests
import hmac
import hashlib
import base64
import time
from docx import Document
from io import BytesIO
import os

# =======================================
# YOUR REAL ConvertAPI SECRET KEY
# =======================================
CONVERTAPI_SECRET = "HESPs6JNV4IDP62WkWpZe3u7ls8KJA38"


# ---------------------------------------
# Create the required ConvertAPI token
# ---------------------------------------
def generate_token(secret):
    expiry = int(time.time()) + 3600  # valid for 1 hour
    message = str(expiry).encode("utf-8")
    signature = hmac.new(secret.encode("utf-8"), message, hashlib.sha256).digest()
    token = base64.urlsafe_b64encode(signature).decode("utf-8")
    return f"{token}-{expiry}"


# ---------------------------------------
# DOCX → PDF conversion function
# ---------------------------------------
def convert_docx_to_pdf(docx_path):

    token = generate_token(CONVERTAPI_SECRET)

    url = f"https://v2.convertapi.com/convert/docx/to/pdf?Token={CONVERTAPI_SECRET}"

    print("Uploading DOCX to ConvertAPI...")
    print("URL:", url)

    with open(docx_path, "rb") as f:
        r = requests.post(url, files={"file": f})

    print("RAW RESPONSE:", r.text)

    if r.status_code != 200:
        raise Exception("API ERROR: " + r.text)

    data = r.json()

    if "Files" not in data:
        raise Exception("ConvertAPI Error: " + r.text)

    pdf_url = data["Files"][0]["Url"]

    print("Downloading PDF from:", pdf_url)

    pdf_bytes = requests.get(pdf_url).content
    return pdf_bytes


# ---------------------------------------
# MAIN TEST SCRIPT
# ---------------------------------------
if __name__ == "__main__":

    DOCX_FILE = "test.docx"

    # Create a simple DOCX if not exists
    if not os.path.exists(DOCX_FILE):
        doc = Document()
        doc.add_heading("ConvertAPI Test", level=1)
        doc.add_paragraph("This PDF was created using ConvertAPI + Python.")
        doc.save(DOCX_FILE)
        print("Created test.docx")

    print("Starting conversion...")
    pdf_data = convert_docx_to_pdf(DOCX_FILE)

    with open("test.pdf", "wb") as f:
        f.write(pdf_data)

    print("🎉 SUCCESS! PDF saved as test.pdf")
